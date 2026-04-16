from docx import Document

doc = Document()

doc.add_heading('Báo cáo chi tiết dự án', level=1)

doc.add_paragraph('Ngày: 2026-04-13')

doc.add_paragraph('1. Tổng quan dự án: Hệ thống truy xuất nguồn gốc nông sản gồm backend (API), blockchain (smart contracts), web (React) và mobile (Flutter).')

doc.add_paragraph('2. Cấu trúc thư mục chính: api/, app/, blockchain/, web/, uploads/, build/, src/.')

doc.add_paragraph('3. Backend (api): agri-trace-api, sử dụng Express, Mongoose, Ethers, JWT, Multer, PDF/Excel generation.')

doc.add_paragraph('4. Blockchain: Hardhat, contracts in contracts/, deploy scripts, typechain-types.')

doc.add_paragraph('5. Web: React app with Axios, QR display, routing; proxy to backend in development.')

doc.add_paragraph('6. Mobile: Flutter app in app/ supporting QR scanning and user interactions.')

doc.add_paragraph('7. Tệp cấu hình: .env, Dockerfile per-service, api/abi/Traceability.json.')

doc.add_paragraph('8. Hướng dẫn chạy: npm install & npm run dev/build for each module; blockchain deploy via Hardhat.')

doc.add_paragraph('9. Ghi chú bảo mật: Không commit bí mật, quản lý private keys, backup MongoDB.')

doc.add_paragraph('Kết luận: Báo cáo đã được tạo và lưu trong Project_Report.rtf; đây là phiên bản DOCX tương đương.')

output = 'Project_Report_fixed.docx'
doc.save(output)
print(f'Wrote {output}')
