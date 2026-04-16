from docx import Document
from docx.shared import Pt

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

doc = Document()

doc.styles['Normal'].font.name = 'Times New Roman'

doc.add_heading('Báo cáo chi tiết: Ứng dụng truy xuất nguồn gốc nông sản sử dụng Blockchain', level=1)
add_paragraph(doc, 'Ngày: 2026-04-13')

doc.add_heading('Tóm tắt', level=2)
add_paragraph(doc, 'Trong bối cảnh chuyển đổi số nông nghiệp, dự án này phát triển một hệ thống truy xuất nguồn gốc nông sản toàn diện kết hợp công nghệ blockchain để tăng tính minh bạch và không thể chối cãi của dữ liệu. Hệ thống gồm: backend API (Node.js/TypeScript), smart contracts (Hardhat/Ethers), frontend web (React), và ứng dụng di động (Flutter). Tính năng nổi bật: ghi nhận, lưu trữ và xác thực thông tin lô hàng (batch), phát sinh mã QR, truy vấn lịch sử phân phối, và xuất báo cáo thống kê cho người quản lý.' )

doc.add_heading('Từ khóa', level=2)
add_paragraph(doc, 'Blockchain, Smart Contract, Traceability, QR code, MongoDB, Hardhat, Ethers, React, Flutter')

doc.add_heading('I. Giới thiệu', level=2)
add_paragraph(doc, 'Vấn đề: chuỗi cung ứng nông sản thiếu minh bạch, gây khó khăn cho việc truy xuất nguồn gốc khi xảy ra vấn đề chất lượng hoặc gian lận. Mục tiêu: xây dựng hệ thống số hóa và phi tập trung một phần (blockchain) để lưu trữ các sự kiện quan trọng liên quan tới sản phẩm, kết hợp hệ thống cơ sở dữ liệu truy vấn nhanh cho ứng dụng web/mobile.')

doc.add_heading('II. Phạm vi và đóng góp', level=2)
add_paragraph(doc, '1) Thiết kế kiến trúc microservices cho hệ thống truy xuất nguồn gốc.')
add_paragraph(doc, '2) Triển khai smart contract lưu sự kiện trọng yếu (tạo batch, chuyển giao, chứng nhận).')
add_paragraph(doc, '3) Tích hợp backend với blockchain qua Ethers để ghi và đọc dữ liệu on-chain/on-db.')
add_paragraph(doc, '4) Cung cấp giao diện web và mobile cho người dùng (nông dân, doanh nghiệp, người tiêu dùng, quản lý).')

doc.add_heading('III. Kiến trúc hệ thống', level=2)
add_paragraph(doc, 'Hệ thống gồm các thành phần chính:')
add_paragraph(doc, '- Smart Contracts: định nghĩa event và hàm để ghi trạng thái batch lên blockchain (phi tập trung).')
add_paragraph(doc, '- Backend API (Node.js/TypeScript): nhận request từ frontend, xác thực, lưu bản sao trong MongoDB, gọi smart contract qua Ethers để ghi sự kiện on-chain.')
add_paragraph(doc, '- Database: MongoDB lưu bản ghi chi tiết, index để truy vấn nhanh; on-chain lưu bằng chứng không thể thay đổi (hash, event).')
add_paragraph(doc, '- Frontend Web (React): Dashboard, quản lý batch, quét QR, báo cáo, lịch sử.')
add_paragraph(doc, '- Mobile App (Flutter): quét QR, xem thông tin truy xuất, nhận thông báo chứng nhận.')

doc.add_heading('IV. Smart Contracts (Chi tiết)', level=2)
add_paragraph(doc, 'Yêu cầu: smart contract phải lưu các sự kiện quan trọng như `CreateBatch`, `TransferBatch`, `CertifyBatch` cùng với timestamp và hash dữ liệu. Contract chứa mappings để tra cứu nhanh theo batchId và emit events để backend/monitor lắng nghe.')
add_paragraph(doc, 'Thiết kế hàm chính:')
add_paragraph(doc, '- createBatch(bytes32 batchHash, string metadataURI) -> emits CreateBatch(batchId, batchHash, owner)')
add_paragraph(doc, '- transferBatch(uint256 batchId, address to) -> emits TransferBatch(batchId, from, to)')
add_paragraph(doc, '- certifyBatch(uint256 batchId, string certificateURI) -> emits CertifyBatch(batchId, certificateURI)')

doc.add_heading('V. Backend & Tích hợp blockchain', level=2)
add_paragraph(doc, 'Backend chịu trách nhiệm: quản lý người dùng (JWT), lưu trữ bản sao dữ liệu trong MongoDB, xử lý file upload (chứng nhận PDF/Excel), tạo QR code và gọi contract thông qua thư viện `ethers`. Ngoài ra backend duy trì một service lắng nghe events on-chain để cập nhật trạng thái trong DB khi có transaction thành công.')

doc.add_heading('VI. Mã QR & Quy trình truy xuất', level=2)
add_paragraph(doc, 'Mỗi `batch` khi tạo sẽ có một mã QR chứa URL/ID; khi quét, frontend gọi API để trả về: thông tin batch, lịch sử chuyển giao (on-chain events), chứng nhận liên quan, các ảnh/ tài liệu liên quan. Để đảm bảo tính xác thực, API cung cấp hash on-chain hoặc link tới transaction trên explorer.')

doc.add_heading('VII. Giao diện người dùng', level=2)
add_paragraph(doc, 'Trang quản trị/giám sát: dashboard tổng hợp số lượng batch, trạng thái chứng nhận, số giao dịch on-chain, biểu đồ thời gian. Trang chi tiết batch: thông tin chi tiết, lịch sử event (kèm txHash), file chứng nhận, nút xuất Excel/PDF. Mobile: quét QR và hiển thị thông tin truy xuất cho người tiêu dùng.')

doc.add_heading('VIII. Bảo mật & Quyền riêng tư', level=2)
add_paragraph(doc, '- Không lưu thông tin nhạy cảm (mật khẩu, private keys) trong repo; private keys dùng để deploy/owner contract lưu trong vault/secret manager.')
add_paragraph(doc, '- JWT cho authentication; RBAC cho quyền quản lý lớp/đơn vị.')
add_paragraph(doc, '- Sử dụng HTTPS/TLS cho giao tiếp giữa client-server; validate và sanitize tất cả dữ liệu upload.')

doc.add_heading('IX. Triển khai', level=2)
add_paragraph(doc, 'Đề xuất triển khai:')
add_paragraph(doc, '- Sử dụng Docker cho từng service (api, frontend, mongodb, node blockchain nếu cần).')
add_paragraph(doc, '- Sử dụng Hardhat để compile và deploy contract; triển khai on testnet (Sepolia) hoặc private chain tùy yêu cầu.')
add_paragraph(doc, '- Sử dụng CI/CD pipeline để build/test và deploy. Khuyến nghị backup periodic cho MongoDB và archive các transaction logs.')

doc.add_heading('X. Dữ liệu và báo cáo', level=2)
add_paragraph(doc, 'Các báo cáo bao gồm: thống kê số batch theo thời gian, tỉ lệ batch có chứng nhận, phân bố nguồn gốc theo vùng, lịch sử chuyển giao. Hỗ trợ export CSV/Excel và PDF để lưu trữ, in ấn.')

doc.add_heading('XI. Kết luận và hướng phát triển', level=2)
add_paragraph(doc, 'Hệ thống truy xuất nguồn gốc nông sản kết hợp blockchain giúp nâng cao tính minh bạch, truy vết và chứng thực dữ liệu. Trong tương lai có thể mở rộng:')
add_paragraph(doc, '- Tích hợp IoT để ghi tự động cảm biến (nhiệt độ, độ ẩm) vào chuỗi dữ liệu.')
add_paragraph(doc, '- Sử dụng zk-SNARKs hoặc giải pháp privacy-preserving nếu cần ẩn thông tin nhạy cảm.')
add_paragraph(doc, '- Tích hợp với sàn thương mại điện tử để minh bạch nguồn gốc tới người tiêu dùng.')

doc.add_heading('Tài liệu tham khảo', level=2)
add_paragraph(doc, '[1] Hardhat documentation - https://hardhat.org')
add_paragraph(doc, '[2] Ethers.js documentation - https://docs.ethers.org')
add_paragraph(doc, '[3] MongoDB documentation - https://www.mongodb.com')
add_paragraph(doc, '[4] React documentation - https://reactjs.org')

output = 'Traceability_Project_Report.docx'
doc.save(output)
print(f'Wrote {output}')
